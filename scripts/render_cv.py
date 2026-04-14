#!/usr/bin/env python3
"""
render_cv.py — Convert tailored CV markdown to formatted DOCX
Usage: python3 render_cv.py <input_markdown.md> <output.docx>

Designed for JobSignal Engine Workflow 3 (The Tailor).
Called by n8n Execute Command node after AI tailoring.

The markdown must follow this structure (sections identified by ALL-CAPS headers):
- Name (first line)
- Contact line (second line)
- Professional summary paragraph
- TECHNICAL SKILLS section (category: items format)
- Experience sections (with bullet points)
- PRIOR EXPERIENCE / EDUCATION sections
"""

import sys
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


def parse_markdown(md_text):
    """Parse the tailored CV markdown into structured sections."""
    lines = md_text.strip().split('\n')
    
    # Clean up lines
    lines = [line.rstrip() for line in lines]
    
    # Remove markdown formatting artifacts
    cleaned = []
    for line in lines:
        # Remove markdown bold markers
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        # Remove markdown heading markers
        line = re.sub(r'^#{1,4}\s+', '', line)
        # Remove horizontal rules
        if re.match(r'^-{3,}$', line.strip()):
            continue
        cleaned.append(line)
    
    return cleaned


def is_section_header(line):
    """Check if a line is an ALL-CAPS section header."""
    stripped = line.strip()
    if not stripped:
        return False
    # Must be mostly uppercase letters and common punctuation
    alpha_chars = re.sub(r'[^a-zA-Z]', '', stripped)
    if not alpha_chars:
        return False
    return (alpha_chars == alpha_chars.upper() 
            and len(alpha_chars) >= 4
            and not stripped.startswith('•')
            and not stripped.startswith('-'))


def is_bullet(line):
    """Check if a line is a bullet point."""
    stripped = line.strip()
    return (stripped.startswith('• ') or 
            stripped.startswith('- ') or 
            stripped.startswith('* '))


def get_bullet_text(line):
    """Extract text from a bullet point line."""
    stripped = line.strip()
    if stripped.startswith('• '):
        return stripped[2:]
    if stripped.startswith('- ') or stripped.startswith('* '):
        return stripped[2:]
    return stripped


def is_skill_line(line):
    """Check if a line is a skills category (e.g., 'CRM & GTM: ...')."""
    stripped = line.strip()
    # Pattern: Category name followed by colon and items
    return bool(re.match(r'^[A-Za-z&\s]+:\s+.+', stripped)) and not stripped.startswith('http')


def is_job_title_line(line):
    """Check if line is a job title/role (bold in original, contains | for location/date)."""
    stripped = line.strip()
    return '|' in stripped and ('–' in stripped or '-' in stripped or '20' in stripped)


def is_project_header(line):
    """Check if line is a project sub-header (descriptive name of a project)."""
    stripped = line.strip()
    if not stripped or is_bullet(stripped) or is_section_header(stripped):
        return False
    if is_job_title_line(stripped) or is_skill_line(stripped):
        return False
    # Project headers are typically title-cased lines that aren't too long
    # and don't start with bullets
    # They often contain em-dashes or parenthetical info
    if ('—' in stripped or '(' in stripped) and len(stripped) > 20 and len(stripped) < 200:
        return True
    return False


def build_docx(lines, output_path):
    """Build the DOCX document matching Rizwan's CV format."""
    doc = Document()
    
    # Page setup - US Letter with tight margins for single-page CV
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    
    # Default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(1)
    paragraph_format.line_spacing = Pt(12)
    
    # Track state
    i = 0
    name_done = False
    contact_done = False
    summary_lines = []
    in_summary = False
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            # If we were collecting summary lines, flush them
            if in_summary and summary_lines:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf = p.paragraph_format
                pf.space_before = Pt(2)
                pf.space_after = Pt(4)
                pf.line_spacing = Pt(12)
                run = p.add_run(' '.join(summary_lines))
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
                summary_lines = []
                in_summary = False
            i += 1
            continue
        
        # First non-empty line = Name
        if not name_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(1)
            run = p.add_run(stripped)
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0D)
            name_done = True
            i += 1
            continue
        
        # Second non-empty line = Contact info
        if name_done and not contact_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(4)
            run = p.add_run(stripped)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            contact_done = True
            in_summary = True
            i += 1
            continue
        
        # Collect summary lines (between contact and first section header)
        if in_summary and not is_section_header(stripped):
            summary_lines.append(stripped)
            i += 1
            continue
        
        # Flush any remaining summary
        if in_summary and summary_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf = p.paragraph_format
            pf.space_before = Pt(2)
            pf.space_after = Pt(4)
            pf.line_spacing = Pt(12)
            run = p.add_run(' '.join(summary_lines))
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            summary_lines = []
            in_summary = False
        
        # Section headers (ALL CAPS)
        if is_section_header(stripped):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(6)
            pf.space_after = Pt(2)
            # Add bottom border
            pBdr = p._p.get_or_add_pPr()
            bdr = pBdr.makeelement(qn('w:pBdr'), {})
            bottom = bdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '4',
                qn('w:space'): '1',
                qn('w:color'): '333333'
            })
            bdr.append(bottom)
            pBdr.append(bdr)
            
            run = p.add_run(stripped)
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x56, 0xAB, 0xB5)
            i += 1
            continue
        
        # Skills lines (Category: items)
        if is_skill_line(stripped):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.line_spacing = Pt(12)
            
            # Split at first colon
            colon_idx = stripped.index(':')
            category = stripped[:colon_idx + 1]
            items = stripped[colon_idx + 1:]
            
            run_cat = p.add_run(category)
            run_cat.font.size = Pt(9.5)
            run_cat.font.bold = True
            run_cat.font.name = 'Calibri'
            
            run_items = p.add_run(items)
            run_items.font.size = Pt(9.5)
            run_items.font.name = 'Calibri'
            i += 1
            continue
        
        # Bullet points
        if is_bullet(stripped):
            bullet_text = get_bullet_text(stripped)
            
            # Check for continuation lines (indented, not a bullet)
            while i + 1 < len(lines):
                next_line = lines[i + 1]
                next_stripped = next_line.strip()
                if (next_stripped 
                    and not is_bullet(next_stripped) 
                    and not is_section_header(next_stripped)
                    and not is_job_title_line(next_stripped)
                    and not is_project_header(next_stripped)
                    and (next_line.startswith('    ') or next_line.startswith('\t'))):
                    bullet_text += ' ' + next_stripped
                    i += 1
                else:
                    break
            
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(1)
            pf.space_after = Pt(1)
            pf.line_spacing = Pt(11.5)
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.15)
            
            run = p.add_run('• ' + bullet_text)
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            i += 1
            continue
        
        # Job title / role lines (contain | separators with dates)
        if is_job_title_line(stripped):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(1)
            pf.line_spacing = Pt(12)
            
            run = p.add_run(stripped)
            run.font.size = Pt(9.5)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            i += 1
            continue
        
        # Project headers (bold descriptive names)
        if is_project_header(stripped):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_before = Pt(4)
            pf.space_after = Pt(1)
            pf.line_spacing = Pt(12)
            
            run = p.add_run(stripped)
            run.font.size = Pt(9.5)
            run.font.bold = True
            run.font.name = 'Calibri'
            i += 1
            continue
        
        # Default: regular paragraph (catch-all for descriptions, etc.)
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(1)
        pf.space_after = Pt(1)
        pf.line_spacing = Pt(12)
        
        run = p.add_run(stripped)
        run.font.size = Pt(9.5)
        run.font.name = 'Calibri'
        i += 1
    
    doc.save(output_path)
    return output_path


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 render_cv.py <input.md> <output.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    lines = parse_markdown(md_text)
    build_docx(lines, output_path)
    print(f"CV generated: {output_path}")


if __name__ == '__main__':
    main()
